import io
import os
import re
import requests
from docx import Document
from dotenv import dotenv_values

env = dotenv_values("/app/frontend/.env")
API = env["REACT_APP_BACKEND_URL"].rstrip("/") + "/api"
r = requests.get(f"{API}/public/contract/download", params={"kode": "CLM-20260817-8U8U", "kontak": "082112223333"})
doc = Document(io.BytesIO(r.content))
lines = [p.text for p in doc.paragraphs]
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            lines.append(c.text)
print("--- lines with underscores ---")
for l in lines:
    if "__" in l:
        print(repr(l))
print("--- lines with braces ---")
for l in lines:
    if "{" in l or "}" in l:
        print(repr(l))
print("--- NIK / branch lines ---")
for l in lines:
    if re.search(r"NIK|NIB|KOTANGSEL|Alamat", l, re.I):
        print(repr(l[:300]))
