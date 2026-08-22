"""E2E: create order via public API -> admin approve with NIK -> verify persistence + contract NIK.

Run explicitly (mutates data): pytest tests/test_nik_approval.py
"""

import io
import json
from datetime import datetime
import os
import re
from pathlib import Path

import pytest
import requests
from docx import Document
from dotenv import dotenv_values

frontend_env = dotenv_values("/app/frontend/.env")
BASE_URL = (os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")).rstrip("/")
API = f"{BASE_URL}/api"
backend_env = dotenv_values("/app/backend/.env")
ADMIN_EMAIL = os.environ.get("TEST_ADMIN_EMAIL") or backend_env.get("TEST_ADMIN_EMAIL")
ADMIN_PASSWORD = os.environ.get("TEST_ADMIN_PASSWORD") or backend_env.get("TEST_ADMIN_PASSWORD")
KTP = Path("/app/tests/assets/ktp_test.png")
NIK = "3276019988776655"
SUFFIX = datetime.now().strftime("%H%M%S")


@pytest.fixture(scope="module")
def admin():
    s = requests.Session()
    r = s.post(f"{API}/auth/login", json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD})
    assert r.status_code == 200, r.text[:200]
    s.headers.update({"Authorization": f"Bearer {r.json()['access_token']}"})
    return s


def _payload(tariff_id: str) -> dict:
    return {
        "nama": "TEST NIK Gate",
        "email": f"test_nikgate_{SUFFIX}@example.com",
        "no_hp": f"0812{SUFFIX}99",
        "alamat_ktp": "Jl. KTP Test No. 9, Tangerang Selatan",
        "provinsi": "BANTEN",
        "kota_kab": "KOTA TANGERANG SELATAN",
        "kecamatan": "PONDOK AREN",
        "kelurahan": "PONDOK BETUNG",
        "detail_alamat": "Jl. Test No. 9 RT01",
        "status_hunian": "Rumah",
        "jenis_ruangan": "Kamar",
        "tanggal_mulai": "2026-09-01",
        "durasi_sewa": 3,
        "catatan": "TEST nik gate",
        "nama_pj_lokasi": "Pak RT Test",
        "no_hp_pj_lokasi": "081222233344",
        "data_consent": True,
        "items": [{"tariff_id": tariff_id, "quantity": 1}],
    }


def _submit_rental(payload: dict) -> str:
    r = requests.post(
        f"{API}/public/rentals",
        data={"payload": json.dumps(payload)},
        files={"ktp": ("ktp.png", KTP.read_bytes(), "image/png")},
    )
    if r.status_code == 429:
        pytest.skip("rate limited (5 submissions/hour per IP)")
    assert r.status_code == 200, r.text[:300]
    return r.json()["kode"]


def _find_order(admin, kode: str) -> dict:
    orders = admin.get(f"{API}/admin/orders", params={"q": kode}).json()
    orders = orders if isinstance(orders, list) else orders.get("items", [])
    return next(o for o in orders if o["kode"] == kode)


def _docx_text(content: bytes) -> str:
    d = Document(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_verify_with_nik_persists_and_renders_in_contract(admin):
    tariffs = requests.get(f"{API}/public/tariffs").json()
    assert tariffs
    payload = _payload(tariffs[0]["id"])
    kode = _submit_rental(payload)
    order = _find_order(admin, kode)

    bad = admin.post(f"{API}/admin/orders/{order['id']}/verify", json={"hasil": "approved"})
    assert bad.status_code == 400 and "NIK" in bad.json()["detail"]

    ok = admin.post(f"{API}/admin/orders/{order['id']}/verify", json={"hasil": "approved", "nik": NIK, "catatan": "TEST"})
    assert ok.status_code == 200, ok.text[:300]
    assert ok.json()["status"] == "verified"

    detail = admin.get(f"{API}/admin/orders/{order['id']}").json()
    assert detail["customer"]["nik"] == NIK
    assert detail["order"]["status"] == "verified"
    assert detail["order"]["contract_status"] == "pending"
    assert detail["contract"] is not None

    doc = requests.get(f"{API}/public/contract/download", params={"kode": kode, "kontak": payload["no_hp"]})
    assert doc.status_code == 200
    text = _docx_text(doc.content)
    assert NIK in text, "NIK not rendered in contract"
    assert "KOTANGSEL" in text and "1701260006922" in text
    assert not re.findall(r"\{[^}\n]{2,60}\}", text)
