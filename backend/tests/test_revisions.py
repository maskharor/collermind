"""Tests for latest revision features (iteration 4).

Covers:
- Public contract DOCX download (template v02, branch/NIB mapping, no leftover placeholders)
- Public + admin invoice DOCX download
- NIK gate on admin verify (negative cases only; approval flow in test_nik_approval)
- Schedule request rules: delivery H+3 minimum, installation jam-only / H+1
- full_access payload contains installation_date / delivered_at keys
"""

import io
import os
import re

import pytest
import requests
from docx import Document
from dotenv import dotenv_values

frontend_env = dotenv_values("/app/frontend/.env")
BASE_URL = (os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")).rstrip("/")
API = f"{BASE_URL}/api"
backend_env = dotenv_values("/app/backend/.env")

ADMIN_EMAIL = os.environ.get("TEST_ADMIN_EMAIL") or backend_env.get("TEST_ADMIN_EMAIL")
ADMIN_PASSWORD = os.environ.get("TEST_ADMIN_PASSWORD") or backend_env.get("TEST_ADMIN_PASSWORD")

E2E_KODE = "CLM-20260817-8U8U"
E2E_KONTAK = "082112223333"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

PLACEHOLDER_RE = re.compile(r"\{[^}\n]{2,60}\}")


@pytest.fixture(scope="module")
def admin():
    s = requests.Session()
    r = s.post(f"{API}/auth/login", json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD})
    if r.status_code != 200:
        pytest.fail(f"Admin login failed: {r.status_code} {r.text[:200]}")
    s.headers.update({"Authorization": f"Bearer {r.json()['access_token']}"})
    return s


def docx_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


# ---------------- Contract DOCX ----------------

class TestContractDownload:
    def test_contract_download_docx_no_placeholder(self):
        r = requests.get(f"{API}/public/contract/download", params={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        assert r.status_code == 200, r.text[:300]
        assert r.headers["content-type"].startswith(DOCX_MIME)
        assert r.content[:2] == b"PK"
        text = docx_text(r.content)
        assert "Dewi Lestari" in text
        left = PLACEHOLDER_RE.findall(text)
        assert not left, f"leftover placeholders: {left[:10]}"
        # only the signature line may keep an underscore blank
        blanks = [ln for ln in text.split("\n") if "__" in ln and "(" not in ln]
        assert not blanks, f"leftover underscore blanks: {blanks[:5]}"

    def test_contract_branch_mapping(self):
        r = requests.get(f"{API}/public/contract/download", params={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        assert r.status_code == 200
        text = docx_text(r.content)
        # order alamat mentions Tangerang Selatan -> KOTANGSEL / NIB 1701260006922
        assert "KOTANGSEL" in text, text[:600]
        assert "1701260006922" in text

    def test_contract_nik_dash_when_missing(self):
        r = requests.get(f"{API}/public/contract/download", params={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        text = docx_text(r.content)
        assert "NIK penyewa" not in text
        assert "Kependudukan (NIK) Nomor -" in text, "NIK fallback '-' not rendered for order without NIK"

    def test_contract_wrong_kontak_403(self):
        r = requests.get(f"{API}/public/contract/download", params={"kode": E2E_KODE, "kontak": "089999999999"})
        assert r.status_code == 403


# ---------------- Invoice DOCX ----------------

class TestInvoiceDownload:
    def test_public_invoice_download(self):
        r = requests.get(f"{API}/public/invoice/download", params={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        assert r.status_code == 200, r.text[:300]
        assert r.headers["content-type"].startswith(DOCX_MIME)
        text = docx_text(r.content)
        assert "Dewi Lestari" in text
        left = PLACEHOLDER_RE.findall(text)
        assert not left, f"leftover placeholders: {left[:10]}"
        assert "INV-CLM-20260817-8U8U" in text

    def test_public_invoice_download_specific_id(self, admin):
        acc = requests.post(f"{API}/public/access", json={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        assert acc.status_code == 200
        invoices = acc.json()["invoices"]
        assert invoices
        inv_id = invoices[0]["id"]
        r = requests.get(f"{API}/public/invoice/download", params={"kode": E2E_KODE, "kontak": E2E_KONTAK, "invoice_id": inv_id})
        assert r.status_code == 200
        assert r.content[:2] == b"PK"

    def test_admin_invoice_download(self, admin):
        acc = requests.post(f"{API}/public/access", json={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        inv_id = acc.json()["invoice"]["id"]
        r = admin.get(f"{API}/admin/invoices/{inv_id}/download")
        assert r.status_code == 200, r.text[:300]
        assert r.headers["content-type"].startswith(DOCX_MIME)
        assert "Dewi Lestari" in docx_text(r.content)

    def test_admin_invoice_download_requires_auth(self):
        r = requests.get(f"{API}/admin/invoices/none/download")
        assert r.status_code in (401, 403)

    def test_admin_invoice_download_404(self, admin):
        r = admin.get(f"{API}/admin/invoices/does-not-exist/download")
        assert r.status_code == 404


# ---------------- NIK gate (negative only) ----------------

class TestNikGate:
    def _pending_order(self, admin):
        r = admin.get(f"{API}/admin/orders", params={"status": "pending"})
        assert r.status_code == 200
        payload = r.json()
        items = payload if isinstance(payload, list) else payload.get("items", [])
        items = [o for o in items if o.get("status") == "pending"]
        return items[0] if items else None

    def test_verify_approved_without_nik_rejected(self, admin):
        order = self._pending_order(admin)
        if not order:
            pytest.skip("no pending order available")
        r = admin.post(f"{API}/admin/orders/{order['id']}/verify", json={"hasil": "approved", "catatan": "TEST"})
        assert r.status_code == 400, r.text[:300]
        assert "NIK" in r.json().get("detail", "")

    def test_verify_approved_with_short_nik_rejected(self, admin):
        order = self._pending_order(admin)
        if not order:
            pytest.skip("no pending order available")
        r = admin.post(f"{API}/admin/orders/{order['id']}/verify", json={"hasil": "approved", "nik": "12345"})
        assert r.status_code == 400
        assert "NIK" in r.json().get("detail", "")

    def test_verify_approved_with_non_digit_nik_rejected(self, admin):
        order = self._pending_order(admin)
        if not order:
            pytest.skip("no pending order available")
        r = admin.post(f"{API}/admin/orders/{order['id']}/verify", json={"hasil": "approved", "nik": "ABCD567890123456"})
        assert r.status_code == 400


# ---------------- Schedule rules ----------------

class TestScheduleRules:
    def _tanggal(self, days):
        from datetime import datetime, timedelta
        from zoneinfo import ZoneInfo
        return (datetime.now(ZoneInfo("Asia/Jakarta")).date() + timedelta(days=days)).isoformat()

    def test_delivery_today_rejected(self):
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "delivery", "tanggal": self._tanggal(0)})
        assert r.status_code == 400, r.text[:300]
        assert "H+3" in r.json().get("detail", ""), r.json()

    def test_delivery_h_plus_2_rejected(self):
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "delivery", "tanggal": self._tanggal(2)})
        assert r.status_code == 400
        assert "H+3" in r.json().get("detail", "")

    def test_delivery_past_date_rejected_422(self):
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "delivery", "tanggal": self._tanggal(-1)})
        assert r.status_code == 422

    def test_delivery_h_plus_5_passes_date_rule(self):
        """H+5 must pass the H+3 rule; it may still fail on order status guard."""
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "delivery", "tanggal": self._tanggal(5)})
        assert r.status_code in (200, 400)
        if r.status_code == 400:
            assert "H+3" not in r.json().get("detail", ""), r.json()

    def test_installation_requires_delivered_status(self):
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "installation", "jam": "09:00"})
        assert r.status_code == 400
        assert "diterima" in r.json().get("detail", "").lower() or "delivered" in r.json().get("detail", "").lower()

    def test_invalid_jenis_rejected(self):
        r = requests.post(f"{API}/public/schedule-request", json={
            "kode": E2E_KODE, "kontak": E2E_KONTAK, "jenis": "bogus", "tanggal": self._tanggal(5)})
        assert r.status_code == 400

    def test_full_access_exposes_installation_date_keys(self):
        r = requests.post(f"{API}/public/access", json={"kode": E2E_KODE, "kontak": E2E_KONTAK})
        assert r.status_code == 200
        data = r.json()
        assert "installation_date" in data and "delivered_at" in data
        if data["delivered_at"]:
            from datetime import date, timedelta
            d = date.fromisoformat(data["delivered_at"][:10])
            assert data["installation_date"] == (d + timedelta(days=1)).isoformat()
