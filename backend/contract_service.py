import io
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Callable
from zoneinfo import ZoneInfo

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

ASSETS_DIR = Path(__file__).parent / "assets"
CONTRACT_TEMPLATE_PATH = ASSETS_DIR / "surat_sewa_v02.docx"
INVOICE_TEMPLATE_PATH = ASSETS_DIR / "invoice.docx"
JAKARTA = ZoneInfo("Asia/Jakarta")

BULAN_ID = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
            "Agustus", "September", "Oktober", "November", "Desember"]

PT_BRANCHES = {
    "EKOKO": {"nib": "1501260007208", "alamat": "Jl. RS. Fatmawati Raya, Blok Aks No. 33-34, Desa/Kelurahan Cipete Utara, Kec. Kebayoran Baru, Kota Adm. Jakarta Selatan, Provinsi DKI Jakarta, Kode Pos: 12150"},
    "FAJARMULIA": {"nib": "1601260025257", "alamat": "Taman Raya Rajeg C-1/6, Desa/Kelurahan Mekarsari, Kec. Rajeg, Kab. Tangerang, Provinsi Banten, Kode Pos: 15545"},
    "KABOGOR": {"nib": "1601260027079", "alamat": "Pesona Prima Cikahuripan 6, Blok A3 No. 10, Desa/Kelurahan Cikahuripan, Kec. Klapanunggal, Kab. Bogor, Provinsi Jawa Barat, Kode Pos: 43366"},
    "KOTANGERANG": {"nib": "1701260006516", "alamat": "Jl. HOS Cokroaminoto, Desa/Kelurahan Sudimara Timur, Kec. Ciledug, Kota Tangerang, Provinsi Banten, Kode Pos: 15151"},
    "KOTANGSEL": {"nib": "1701260006922", "alamat": "Jl. Pd. Betung Raya, Desa/Kelurahan Pondok Betung, Kec. Pondok Aren, Kota Tangerang Selatan, Provinsi Banten, Kode Pos: 15221"},
    "KODEPOK": {"nib": "1701260007361", "alamat": "Jl. Alternatif Cibubur, Desa/Kelurahan Harjamukti, Kec. Cimanggis, Kota Depok, Provinsi Jawa Barat, Kode Pos: 16454"},
    "KOBOGOR": {"nib": "1701260007618", "alamat": "Jl. Raya Pajajaran, Desa/Kelurahan Tegallega, Kec. Bogor Tengah, Kota Bogor, Provinsi Jawa Barat, Kode Pos: 16129"},
    "JAKPUS": {"nib": "1801260007362", "alamat": "Jl. Administrasi Negara, Desa/Kelurahan Bendungan Hilir, Kec. Tanah Abang, Kota Adm. Jakarta Pusat, Provinsi DKI Jakarta, Kode Pos: 10210"},
    "JAKTIM": {"nib": "1801260007755", "alamat": "Jl. Cipinang Jaya Raya, Desa/Kelurahan Cipinang Muara, Kec. Jatinegara, Kota Adm. Jakarta Timur, Provinsi DKI Jakarta, Kode Pos: 13410"},
    "JAKSEL": {"nib": "1801260008712", "alamat": "Jl. Raya Jagakarsa, Desa/Kelurahan Jagakarsa, Kec. Jagakarsa, Kota Adm. Jakarta Selatan, Provinsi DKI Jakarta, Kode Pos: 12610"},
    "JAKUT": {"nib": "1801260013155", "alamat": "Jl. Cibanteng, Koja, Desa/Kelurahan Kebon Bawang, Kec. Tanjung Priok, Kota Adm. Jakarta Utara, Provinsi DKI Jakarta, Kode Pos: 14320"},
    "JAKBAR": {"nib": "1801260013785", "alamat": "Jl. Strategi III, Desa/Kelurahan Joglo, Kec. Kembangan, Kota Adm. Jakarta Barat, Provinsi DKI Jakarta, Kode Pos: 11640"},
    "KABEKASI": {"nib": "1801260015258", "alamat": "Jl. H. Jampang, Desa/Kelurahan Jatimulya, Kec. Tambun Selatan, Kab. Bekasi, Provinsi Jawa Barat, Kode Pos: 17510"},
    "KOBEKASI": {"nib": "1801260016316", "alamat": "Cluster Taman Jati Kramat Indah Jl. Jatikramat II Blok A. 111, Desa/Kelurahan Jatikramat, Kec. Jatiasih, Kota Bekasi, Provinsi Jawa Barat, Kode Pos: 17421"},
}

BranchMatcher = Callable[[str, str], bool]
_BRANCH_RULES: list[tuple[BranchMatcher, str]] = [
    (lambda kota, text: "tangerang selatan" in text or "tangsel" in text, "KOTANGSEL"),
    (lambda kota, text: "kabupaten tangerang" in kota, "FAJARMULIA"),
    (lambda kota, text: "tangerang" in text, "KOTANGERANG"),
    (lambda kota, text: "depok" in text, "KODEPOK"),
    (lambda kota, text: "kota bogor" in kota, "KOBOGOR"),
    (lambda kota, text: "bogor" in text, "KABOGOR"),
    (lambda kota, text: "kota bekasi" in kota, "KOBEKASI"),
    (lambda kota, text: "bekasi" in text, "KABEKASI"),
    (lambda kota, text: "jakarta pusat" in text, "JAKPUS"),
    (lambda kota, text: "jakarta timur" in text, "JAKTIM"),
    (lambda kota, text: "jakarta utara" in text, "JAKUT"),
    (lambda kota, text: "jakarta barat" in text, "JAKBAR"),
    (lambda kota, text: "jakarta selatan" in text, "JAKSEL"),
]


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").lower().replace("adm.", "").replace("kab.", "kabupaten")).strip()


def detect_pt_branch(kota_kab: str = "", alamat: str = "") -> dict:
    kota = _norm(kota_kab)
    text = f"{kota} {_norm(alamat)}"
    code = "EKOKO"
    for matches, branch_code in _BRANCH_RULES:
        if matches(kota, text):
            code = branch_code
            break
    return {"code": code, **PT_BRANCHES[code]}


def _rp(n) -> str:
    return "Rp" + f"{int(n or 0):,}".replace(",", ".")


def _num(n) -> str:
    return f"{int(n or 0):,}".replace(",", ".")


_SATUAN = ["", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "delapan", "sembilan", "sepuluh", "sebelas"]


def terbilang(n) -> str:
    n = int(n or 0)
    if n == 0:
        return "nol"
    if n < 12:
        return _SATUAN[n]
    if n < 20:
        return f"{_SATUAN[n - 10]} belas"
    if n < 100:
        return f"{_SATUAN[n // 10]} puluh {terbilang(n % 10)}".strip()
    if n < 200:
        return f"seratus {terbilang(n - 100)}".strip()
    if n < 1000:
        return f"{_SATUAN[n // 100]} ratus {terbilang(n % 100)}".strip()
    if n < 2000:
        return f"seribu {terbilang(n - 1000)}".strip()
    if n < 1_000_000:
        return f"{terbilang(n // 1000)} ribu {terbilang(n % 1000)}".strip()
    if n < 1_000_000_000:
        return f"{terbilang(n // 1_000_000)} juta {terbilang(n % 1_000_000)}".strip()
    return f"{terbilang(n // 1_000_000_000)} miliar {terbilang(n % 1_000_000_000)}".strip()


def _parse_dt(value) -> datetime:
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if dt.tzinfo is None:
            return dt.replace(tzinfo=ZoneInfo("UTC")).astimezone(JAKARTA)
        return dt.astimezone(JAKARTA)
    except Exception:
        logger.exception("invalid document datetime")
        return datetime.now(JAKARTA)


def _long_date(value) -> str:
    dt = _parse_dt(value)
    return f"tanggal {dt.day} bulan {BULAN_ID[dt.month]} tahun {dt.year}"


def _xml_paragraphs(doc) -> list:
    return doc.element.xpath(".//w:p")


def _p_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.xpath(".//w:t"))


def _clear_highlights(p_el) -> None:
    for hl in p_el.xpath(".//w:highlight"):
        parent = hl.getparent()
        if parent is not None:
            parent.remove(hl)


def _p_set(p_el, value: str) -> None:
    texts = p_el.xpath(".//w:t")
    if not texts:
        return
    texts[0].text = str(value)
    for t in texts[1:]:
        t.text = ""
    _clear_highlights(p_el)


def _replace_placeholders(doc, subs: list[tuple[str, str]]) -> None:
    for p_el in _xml_paragraphs(doc):
        text = _p_text(p_el)
        if not text:
            continue
        new = text
        for pat, rep in subs:
            new = re.sub(pat, str(rep), new, flags=re.IGNORECASE)
        if new != text:
            _p_set(p_el, new)


def _set_paragraph_if_contains(doc, needle: str, value: str) -> None:
    for p_el in _xml_paragraphs(doc):
        if needle.lower() in _p_text(p_el).lower():
            _p_set(p_el, value)


def _set_cell(cell, value, size: int | None = None) -> None:
    paragraphs = cell.paragraphs or [cell.add_paragraph()]
    for idx, p in enumerate(paragraphs):
        if p.runs:
            p.runs[0].text = str(value) if idx == 0 else ""
            for r in p.runs[1:]:
                r.text = ""
        elif idx == 0:
            p.add_run(str(value))
        for r in p.runs:
            r.font.highlight_color = None
            if size:
                r.font.size = Pt(size)


def _unique_cells(cells: list) -> list:
    seen, out = set(), []
    for cell in cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        out.append(cell)
    return out


def generate_contract_docx(order: dict, customer: dict, contract: dict) -> bytes:
    doc = Document(str(CONTRACT_TEMPLATE_PATH))
    branch = detect_pt_branch(customer.get("kota_kab", ""), customer.get("alamat_pemasangan", ""))
    details = order.get("details", [])
    total_unit = sum(d.get("quantity", 0) for d in details)
    items_desc = ", ".join(f"{d['nama']} ({d['quantity']} unit)" for d in details)
    sewa = sum(d.get("harga", 0) * d.get("quantity", 0) for d in details)
    total_awal = (order.get("estimasi") or {}).get("total", sewa + 650000)

    subs = [
        (r"tanggal\s*_+\s*bulan\s*_+\s*tahun\s*_+", _long_date(contract.get("created_at"))),
        (r"\{daerah cabang\}", branch["code"]),
        (r"\{NIB Cabang\}", branch["nib"]),
        (r"\{Alamat lengkap daerah cabang\}", branch["alamat"]),
        (r"\(Nama lengkap\)", customer.get("nama", "-")),
        (r"\{Alamat seusai KTP penyewa\s*\}", customer.get("alamat_ktp", "-")),
        (r"\{NIK penyewa\}", customer.get("nik") or "-"),
        (r"\{no telp\}", customer.get("no_hp", "-")),
        (r"sebanyak\s*_+\s*\(_+\)\s*unit Air AC dengan kapasitas\s*_+\s*PK \(Standart/Inverter\)",
         f"sebanyak {total_unit} ({terbilang(total_unit)}) unit Air AC dengan kapasitas: {items_desc}"),
        (r"\{detail Alamat pemasangan yang di input customer\}", customer.get("alamat_pemasangan", "-")),
        (r"Rp\s*_\s*\{nominal angka\},-\s*\{nominal terbilagn\}", f"{_rp(total_awal)},- ({terbilang(total_awal)} rupiah)"),
        (r"Biaya sewa bulan pertama: Rp_+", f"Biaya sewa bulan pertama: {_rp(sewa)}"),
        (r"jangka waktu 1 \(satu\) bulan kalender", f"jangka waktu {order.get('durasi_sewa', '-')} bulan kalender"),
        (r"selanjutnya sebesar Rp_+", f"selanjutnya sebesar {_rp(sewa)}"),
    ]
    _replace_placeholders(doc, subs)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _invoice_context(order: dict, customer: dict, invoice: dict) -> dict:
    branch = detect_pt_branch(customer.get("kota_kab", ""), customer.get("alamat_pemasangan", ""))
    details = order.get("details", [])
    dt = _parse_dt(invoice.get("issued_at") or invoice.get("created_at"))
    extra_meter = invoice.get("extra_pipa_meter") or 0
    extra_cost = invoice.get("biaya_extra_pipa") or 0
    note = "Pembayaran dilakukan setelah instalasi selesai dan invoice terbit."
    if extra_meter:
        note = f"Extra pipa {extra_meter} m × {_rp(130000)} = {_rp(extra_cost)} (pipa terpakai {invoice.get('total_pipa_meter', '-')} m). " + note
    return {
        "branch": branch,
        "customer_nama": customer.get("nama", "-"),
        "alamat_ktp": customer.get("alamat_ktp", "-"),
        "total_unit": sum(d.get("quantity", 0) for d in details),
        "kapasitas": " / ".join(sorted({d.get("kapasitas", "").replace(" PK", "") for d in details if d.get("kapasitas")})) or "-",
        "variants": " / ".join(sorted({d.get("variant", "Standart") for d in details})) or "Standart",
        "sewa": next((i.get("amount", 0) for i in invoice.get("items", []) if "Sewa bulan pertama" in i.get("label", "")), 0),
        "total": invoice.get("total", 0),
        "dt": dt,
        "invoice_date": dt.strftime("%d/%m/%Y"),
        "nomor": invoice.get("nomor", "-"),
        "rekening": invoice.get("rekening", "-"),
        "note": note,
    }


def _invoice_subs(ctx: dict) -> list[tuple[str, str]]:
    return [
        (r"\{tanggal invoice dd/mm/yyyy\}", ctx["invoice_date"]),
        (r"\{nama lengkap\}", ctx["customer_nama"]),
        (r"\{nama lengkap customer\}", ctx["customer_nama"]),
        (r"\{Alamat sesuai KTP[^}]*\}", ctx["alamat_ktp"]),
        (r"\{kapasistas PK\}", ctx["kapasitas"]),
        (r"\{kapasitas PK\}", ctx["kapasitas"]),
        (r"\{\{?tipe Inverter / Standart\}\}?", ctx["variants"]),
        (r"\{banyak unit\}", str(ctx["total_unit"])),
        (r"\{biaya nominal angka\}", _num(ctx["sewa"])),
        (r"\{total yang dibayar\}", _num(ctx["total"])),
        (r"\{nominal terbilang\}", f"{terbilang(ctx['total'])} rupiah"),
        (r"\{rekening tujuan[^}]*\}", ctx["rekening"]),
        (r"\{[^}]*daerah cabang[^}]*\}", ctx["branch"]["code"]),
        (r"tanggal\s*_+\s*bulan\s*_+\s*tahun\s*_+", f"tanggal {ctx['dt'].day} bulan {BULAN_ID[ctx['dt'].month]} tahun {ctx['dt'].year}"),
    ]


def _fill_invoice_tables(doc, ctx: dict) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    right_cells = _unique_cells([table.rows[0].cells[-1], table.rows[1].cells[-1]])
    if len(right_cells) == 1:
        _set_cell(right_cells[0], f"{ctx['invoice_date']}\n{ctx['nomor']}", size=8)
    elif len(right_cells) > 1:
        _set_cell(right_cells[0], ctx["invoice_date"], size=9)
        _set_cell(right_cells[1], ctx["nomor"], size=8)
    _set_cell(table.rows[1].cells[1], ctx["customer_nama"])
    _set_cell(table.rows[2].cells[1], ctx["alamat_ktp"], size=9)
    _set_cell(table.rows[6].cells[1], f"Pemasangan Unit AC Split {ctx['kapasitas']} PK {ctx['variants']}", size=9)
    _set_cell(table.rows[7].cells[1], f"Sewa AC Baru {ctx['kapasitas']} PK {ctx['variants']}", size=9)
    _set_cell(table.rows[7].cells[3], _num(ctx["sewa"]), size=9)
    _set_cell(table.rows[7].cells[-1], _num(ctx["sewa"]), size=9)
    _set_cell(table.rows[9].cells[1], f"Note\n{ctx['note']}", size=9)
    payment_text = f"Silahkan melakukan pembayaran ke :\n{ctx['rekening']}\nPT Mula Collermind {ctx['branch']['code']}"
    payment_cells = _unique_cells([table.rows[10].cells[1], table.rows[11].cells[1], table.rows[12].cells[1]])
    if payment_cells:
        _set_cell(payment_cells[0], payment_text, size=9)
    _set_cell(table.rows[10].cells[-1], _num(ctx["total"]), size=9)
    terbilang_cells = _unique_cells(table.rows[12].cells[3:])
    if terbilang_cells:
        _set_cell(terbilang_cells[0], f"{terbilang(ctx['total'])} rupiah", size=9)
    if len(doc.tables) > 1:
        bast = doc.tables[1]
        _set_cell(bast.rows[0].cells[1], f"Pemasangan Unit AC Split {ctx['kapasitas']} PK", size=9)
        _set_cell(bast.rows[0].cells[2], f"{ctx['total_unit']} Unit", size=9)
        _set_cell(bast.rows[1].cells[1], f"AC {ctx['kapasitas']} PK {ctx['variants']}", size=9)
        _set_cell(bast.rows[1].cells[2], f"{ctx['total_unit']} Unit", size=9)


def generate_invoice_docx(order: dict, customer: dict, invoice: dict) -> bytes:
    doc = Document(str(INVOICE_TEMPLATE_PATH))
    ctx = _invoice_context(order, customer, invoice)
    _replace_placeholders(doc, _invoice_subs(ctx))
    _set_paragraph_if_contains(doc, "PT Mula Collermind {", f"PT Mula Collermind {ctx['branch']['code']}")
    _set_paragraph_if_contains(doc, "{nama lengkap customer}", f"{ctx['customer_nama']}\t\t\tPT Mula Collermind {ctx['branch']['code']}")
    for needle in ("sesuai alamat  pemasangan", "kota/kabupaten ...}"):
        _set_paragraph_if_contains(doc, needle, "")
    try:
        _fill_invoice_tables(doc, ctx)
    except Exception:
        logger.exception("invoice template table mapping failed")
        raise
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
